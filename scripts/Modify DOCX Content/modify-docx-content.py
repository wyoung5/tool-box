import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_tracked_change(paragraph, text, is_insertion=True, author="Script", date=None):
    """Add a tracked change (insertion or deletion) to a paragraph."""
    run = paragraph.add_run(text)

    # Create the w:ins or w:del element for tracked changes
    if is_insertion:
        run._r.get_or_add_ins(author=author, date=date)
    else:
        run._r.get_or_add_del(author=author, date=date)

def process_docx_files(file1_path, file2_path, file3_path, output_path):
    # Load all three documents
    doc1 = Document(file1_path)
    doc2 = Document(file2_path)
    doc3 = Document(file3_path)

    # Create a new document for the output with change tracking enabled
    new_doc = Document()
    new_doc.settings.track_revisions = True  # Enable change tracking

    # Process paragraphs from file1, with modifications from file2
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        new_para = new_doc.add_paragraph()

        # Copy runs from file1
        for run1 in para1.runs:
            run_text = run1.text.strip()
            if not run_text:
                continue  # Skip empty runs

            # Check if corresponding run in file2 has highlighting
            if para2.runs and run_text == para2.runs[0].text.strip():
                # Apply formatting from file2 if it's highlighted
                if para2.runs[0].font.highlight_color:
                    # Mark as insertion (since we're adding highlighting)
                    add_tracked_change(new_para, run1.text, is_insertion=True)
                    # Apply the highlighting to the run
                    run_new = new_para.runs[-1]
                    run_new.font.highlight_color = para2.runs[0].font.highlight_color
            else:
                # Regular text without tracked changes
                new_para.add_run(run1.text)

    # Add content from file3 as insertions (tracked changes)
    for para in doc3.paragraphs:
        add_tracked_change(new_doc.add_paragraph(), para.text)

    # Save the new document with tracked changes
    new_doc.save(output_path)
    print(f"New document with tracked changes created at: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python docx_combiner_tracked.py <file1.docx> <file2.docx> <file3.docx> <output.docx>")
        sys.exit(1)

    file1_path = sys.argv[1]
    file2_path = sys.argv[2]
    file3_path = sys.argv[3]
    output_path = sys.argv[4]

    process_docx_files(file1_path, file2_path, file3_path, output_path)
