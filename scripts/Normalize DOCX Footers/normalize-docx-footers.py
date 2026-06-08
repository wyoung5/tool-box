from docx import Document

def copy_footer_content(source_footer, target_footer):
    """Copy all paragraphs and runs from one footer to another."""
    # Clear target footer
    for p in target_footer.paragraphs:
        p._element.getparent().remove(p._element)

    # Copy content
    for para in source_footer.paragraphs:
        new_para = target_footer.add_paragraph()
        new_para.style = para.style

        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

def normalize_footers(doc_path, output_path, reference_section_index=0):
    doc = Document(doc_path)

    # Get reference footer
    ref_section = doc.sections[reference_section_index]
    ref_footer = ref_section.footer

    # Normalize each section
    for section in doc.sections:
        # Disable different first / odd-even settings
        section.different_first_page_header_footer = False
        section.even_page_footer.is_linked_to_previous = False

        footer = section.footer
        footer.is_linked_to_previous = False

        # Copy reference footer content
        copy_footer_content(ref_footer, footer)

    # Save result
    doc.save(output_path)
    print(f"Normalized document saved to: {output_path}")


# ✅ Usage
input_file = r"C:\Users\wyoung5\OneDrive - McGill University\CRU\Tools\Normalize DOCX Footers\ICF Main (EN V3 1Jun2026)_ARGX-117-2202_TC.docx"
output_file = "output_normalized.docx"

normalize_footers(input_file, output_file)